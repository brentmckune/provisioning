"""Document ingestion: load policy documents from a directory.

Supports .txt and .md natively; .docx via python-docx and .pdf via pypdf
when those optional dependencies are installed.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

SUPPORTED_EXTENSIONS = {".txt", ".md", ".markdown", ".docx", ".pdf"}

# Documents larger than this are truncated with a warning rather than sent
# whole. ~1.5M characters is roughly 400K tokens — comfortably inside the
# model's 1M-token window even with the knowledge base and instructions.
MAX_CHARS_PER_DOCUMENT = 1_500_000


@dataclass
class PolicyDocument:
    name: str  # filename, used as the policy identifier throughout
    path: Path
    text: str
    truncated: bool = False


class IngestionError(Exception):
    pass


def _read_docx(path: Path) -> str:
    try:
        import docx  # type: ignore
    except ImportError as e:
        raise IngestionError(
            f"Cannot read {path.name}: python-docx is not installed. "
            "Install it with: pip install python-docx"
        ) from e
    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError as e:
        raise IngestionError(
            f"Cannot read {path.name}: pypdf is not installed. "
            "Install it with: pip install pypdf"
        ) from e
    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def load_document(path: Path) -> PolicyDocument:
    ext = path.suffix.lower()
    if ext in {".txt", ".md", ".markdown"}:
        text = path.read_text(encoding="utf-8", errors="replace")
    elif ext == ".docx":
        text = _read_docx(path)
    elif ext == ".pdf":
        text = _read_pdf(path)
    else:
        raise IngestionError(f"Unsupported file type: {path.name}")

    text = text.strip()
    if not text:
        raise IngestionError(f"{path.name} contains no extractable text.")

    truncated = False
    if len(text) > MAX_CHARS_PER_DOCUMENT:
        text = text[:MAX_CHARS_PER_DOCUMENT]
        truncated = True

    return PolicyDocument(name=path.name, path=path, text=text, truncated=truncated)


def load_directory(directory: str | Path) -> list[PolicyDocument]:
    """Load every supported policy document in a directory (non-recursive by default)."""
    dir_path = Path(directory)
    if not dir_path.is_dir():
        raise IngestionError(f"Not a directory: {directory}")

    docs: list[PolicyDocument] = []
    errors: list[str] = []
    for path in sorted(dir_path.rglob("*")):
        if not path.is_file() or path.name.startswith("."):
            continue
        if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        try:
            docs.append(load_document(path))
        except IngestionError as e:
            errors.append(str(e))

    if not docs:
        detail = ("\n" + "\n".join(errors)) if errors else ""
        raise IngestionError(
            f"No readable policy documents found in {directory}. "
            f"Supported types: {', '.join(sorted(SUPPORTED_EXTENSIONS))}{detail}"
        )
    for err in errors:
        print(f"warning: {err}")
    return docs
